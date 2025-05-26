# Text Summarizer App
import streamlit as st
from transformers import pipeline
from docx import Document
import fitz  # PyMuPDF
from io import BytesIO
from time import sleep

# Configure Text-Summarizer-App layout:
st.set_page_config(
    page_title="Text Summarizer App",   # setup title for browser-tab
    page_icon="📄",                     # setup favicon for browser-tab
    layout="wide",                      # setup UI-display-mode to "wide"
    initial_sidebar_state="expanded",   # setup sidebar mode to "expanded"
    menu_items={                        # setup three-dots-menu, top-right corner of app-UI (A Dictionary construct: can have three key:value pairs, with key-names; "Get help": "a valid URL-address", "Report a bug":"a valid URL-address" and "About":"any text")
        "About": "This app summarizes user-uploaded text files (TXT, DOCX, or PDF). It generates a concise summary using the Generative-AI LLM model: facebook/bart-large-cnn. The summary can be downloaded in TXT, DOCX, or PDF format. The app was designed and implemented with the assistance of ChatGPT.",
    }
)

# Define main function
def main_function():      

    # Call function to setup UI
    max_len, min_len = setup_ui()

    # Call function to load summarizer model
    summarizer = summarizer_model()  

    # initialize session_state-flages & variables
    if 'file_uploaded' not in st.session_state:       
        st.session_state.download_triggered = False
        st.session_state.file_uploaded = False
        st.session_state.word_count_uploaded = " "
        st.session_state.word_count_summary = " "
        st.session_state.downloadable_files = {}
        st.session_state.final_summary = " "        
        st.session_state.file = None
        st.session_state.text = " "        

    if not st.session_state.file_uploaded:          
                
        # Call function to upload file
        file = upload_file()
        if file:
            st.session_state.file_uploaded = True
            st.session_state.file = file

        # Check if file has been uploaded/avaialable before proceeding
        if st.session_state.file:
            file = st.session_state.file
            
            # Call function to display and summarize the uploaded text
            final_summary, text, word_count_uploaded, word_count_summary = display_and_summarize(file, summarizer, max_len, min_len)

            # Update session-state variables
            st.session_state.final_summary = final_summary
            st.session_state.text = text
            st.session_state.word_count_uploaded = word_count_uploaded
            st.session_state.word_count_summary = word_count_summary

            # Call function to generate downloadable files of summary            
            downloadable_files = create_downloadable_files(st.session_state.final_summary)
            # update session_state variable
            st.session_state.downloadable_files = downloadable_files

    # Check if downloadable_files have been created, then call function to create buttons and download files
    if st.session_state.downloadable_files:        
        download_buttons(st.session_state.downloadable_files)     
            
    # Check if session_state download_triggered flag is TRUE, then call function to display original-text, summary and word-counts
    if st.session_state.download_triggered:
        # Update variables from session_state 
        final_summary = st.session_state.final_summary
        text = st.session_state.text
        word_count_uploaded = st.session_state.word_count_uploaded
        word_count_summary = st.session_state.word_count_summary       
        
        display_content_word_counts(final_summary, text, word_count_uploaded, word_count_summary)         
        st.session_state.download_triggered = False        # Update session_state flag
    
# Define function to access summarizer model
def summarizer_model():
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")    
    #summarizer = pipeline("summarization", model="facebook/bart-base")
    
    return summarizer

# Define to setup UI
def setup_ui():   
    
    st.header("**Text Summarization Application**")
    st.text("The App generates a concise summary of the uploaded text.")

    # Setup "About-App" button and App intro
    if st.sidebar.button(label="About App"):
        # App intro is coded in html/Markdown-script.
        st.sidebar.markdown(
            '<p style="font-size:12px;">This app summarizes user-uploaded text files (TXT, DOCX, or PDF). It generates a concise summary using the Generative-AI LLM model: facebook/bart-large-cnn. The summary can be downloaded in TXT, DOCX, or PDF format. The app was designed and implemented with the assistance of ChatGPT.</p>',
            unsafe_allow_html=True  # code line required for using html in streamlit.
        )

    # Setup adjustable summary length options.
    st.sidebar.header("Summary Length Adjustment:")
    max_len = st.sidebar.slider("Maximum Length", min_value=50, max_value=500, value=150)
    min_len = st.sidebar.slider("Minimum Length", min_value=20, max_value=100, value=50)
    return max_len, min_len

# Define file upload function
def upload_file():    
    return st.file_uploader("Upload a text file (TXT, DOCX, or PDF) by clicking 'Browse Files' or dragging and dropping a file.", type=["txt", "docx", "pdf"])

# Define summarizer function
def display_and_summarize(file, summarizer, max_len, min_len):    
    # Determine file extension
    if file.name.endswith(".docx"):
        file_extension = "docx"
    elif file.name.endswith(".pdf"):
        file_extension = "pdf"
    else:
        file_extension = "txt"

    # Call function to extract content of uploaded file
    text = extract_text(file, file_extension)

    # If the file is empty, show an error message
    if not text.strip():
        st.error("The uploaded file is empty. Please upload a file with content.")
        return

    # Calculate word count of the uploaded file
    word_count_uploaded = len(text.split())

    # Call function to split the text into chunks
    chunks = chunk_text(text)

    summaries = []  # Initialize an empty list to store chunk summaries
    progress_bar = st.progress(0)  # Initialize progress bar

    try:
        with st.spinner("Processing..."):            
            for i, chunk in enumerate(chunks):
                summary = summarizer(chunk, max_length=max_len, min_length=min_len, do_sample=False)[0]['summary_text']
                summaries.append(summary)

                # Update the progress bar as a percentage of chunks completed
                progress_bar.progress((i + 1) / len(chunks))

        # Reset the progress bar after summarization is complete
        progress_bar.progress(100)

    except Exception as e:
        st.error(f"Error during summarization: {str(e)}")
        progress_bar.empty()  # Clear progress bar in case of an error
        return

    # Combine the chunk summaries into one final summary
    final_summary = " ".join(summaries)

    # Calculate word count of the summary
    word_count_summary = len(final_summary.split())

    # if session_state.download_triggered flag is False, call function to display content and word-counts
    if not st.session_state.download_triggered:        
        display_content_word_counts(final_summary, text, word_count_uploaded, word_count_summary)
        return final_summary, text, word_count_uploaded, word_count_summary

# Define function to extract-text content from different file types (txt, docx, pdf)
def extract_text(file, file_type):    
    if file_type == "txt":
        return file.read().decode("utf-8")
    elif file_type == "docx":
        doc = Document(file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif file_type == "pdf":
        pdf_text = ""
        pdf_doc = fitz.open(stream=file.read(), filetype="pdf")
        for page_num in range(pdf_doc.page_count):
            page = pdf_doc.load_page(page_num)
            pdf_text += page.get_text()
        pdf_doc.close()
        return pdf_text
    else:
        return None

# Define function to split the text into chunks at word boundaries
def chunk_text(text, max_length=1000):   
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0
    for word in words:
        if current_length + len(word) + 1 <= max_length:
            current_chunk.append(word)
            current_length += len(word) + 1
        else:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = len(word) + 1
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks

# Define function to pre-generate downloadable files in txt, docx, pdf formats.
def create_downloadable_files(final_summary):    
    # Create .txt file
    txt_file = BytesIO(final_summary.encode("utf-8"))
    # Create .docx file
    docx_file = BytesIO()
    doc = Document()
    doc.add_paragraph(final_summary)
    doc.save(docx_file)
    docx_file.seek(0)
    # Call function to create .pdf file
    pdf_file = generate_pdf(final_summary)
    
    return {        # return data as a dictionary of dictionaries
        ".txt": {"data": txt_file, "mime": "text/plain", "file_name": "Summary.txt"},
        ".docx": {"data": docx_file, "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "file_name": "Summary.docx"},
        ".pdf": {"data": pdf_file, "mime": "application/pdf", "file_name": "Summary.pdf"},
    }

# Define function to create .pdf file
def generate_pdf(final_summary):
    # Initialize an in-memory PDF
    pdf_file = BytesIO()
    pdf_doc = fitz.open()  # Create a new PDF document

    # PDF page settings
    page_width, page_height = 595, 842  # A4 dimensions in points
    margin = 72  # 1-inch margin
    font_size = 12
    line_spacing = 1.5  # Line spacing factor
    text_width_limit = page_width - 2 * margin  # Usable width

    # Font and text style
    font_name = "helv"  # Helvetica
    text_lines = final_summary.splitlines()  # Split text into lines

    # Create the first page and set starting position
    page = pdf_doc.new_page(width=page_width, height=page_height)
    x, y = margin, margin

    # Iterate through each line of text
    for line in text_lines:
        words = line.split()  # Split line into words
        current_line = ""

        for word in words:
            # Check the width of the line with the current word
            test_line = f"{current_line} {word}".strip()
            test_width = fitz.get_text_length(test_line, fontsize=font_size, fontname=font_name)

            if test_width <= text_width_limit:
                # If the current word fits, add it to the line
                current_line = test_line
            else:
                # Print the current line and move to the next
                page.insert_text((x, y), current_line, fontsize=font_size, fontname=font_name)
                y += font_size * line_spacing  # Move cursor down

                # Start the new line with the current word
                current_line = word

                # Check if we're out of space on the current page
                if y + font_size * line_spacing > page_height - margin:
                    # Start a new page
                    page = pdf_doc.new_page(width=page_width, height=page_height)
                    y = margin

        # Print the last remaining line on the page
        if current_line:
            page.insert_text((x, y), current_line, fontsize=font_size, fontname=font_name)
            y += font_size * line_spacing

        # Check if we're out of space after the last line
        if y + font_size * line_spacing > page_height - margin:
            # Start a new page
            page = pdf_doc.new_page(width=page_width, height=page_height)
            y = margin

    # Save the PDF to memory
    pdf_doc.save(pdf_file)
    pdf_doc.close()
    pdf_file.seek(0)  # Reset file pointer
    return pdf_file

# Define function to display original-text, summary and word counts in two columns
def display_content_word_counts(final_summary, text, word_count_uploaded, word_count_summary):    
    col1, col2 = st.columns(2)
    col1.text_area(f"Uploaded Content (Word Count: {word_count_uploaded})", text, height=500)
    col2.text_area(f"Summary (Word Count: {word_count_summary})", final_summary, height=500)

# Define function to create buttons and download files
def download_buttons(downloadable_files):      
    st.sidebar.text("Download Summary:")
    st.sidebar.download_button(
        label="Download as .txt",
        data=downloadable_files[".txt"]["data"],
        file_name=downloadable_files[".txt"]["file_name"],
        mime=downloadable_files[".txt"]["mime"],
        on_click=trigger_download,
    )
    st.sidebar.download_button(
        label="Download as.docx",
        data=downloadable_files[".docx"]["data"],
        file_name=downloadable_files[".docx"]["file_name"],
        mime=downloadable_files[".docx"]["mime"],
        on_click=trigger_download,
    )
    st.sidebar.download_button(
        label="Download as .pdf",
        data=downloadable_files[".pdf"]["data"],
        file_name=downloadable_files[".pdf"]["file_name"],
        mime=downloadable_files[".pdf"]["mime"],
        on_click=trigger_download,
    ) 
    # If "Upload Text File" button is clicked, then reset session_state and rerun the app from top of the script
    if st.sidebar.button("Upload Text File"): 
        reset_session_state()                  
        st.rerun()
        
# Define function to set session_state.download_triggered flag.
def trigger_download():    
    st.session_state.download_triggered = True
    
# Define function to reset session_state-flags & variables
def reset_session_state():    
    st.session_state.download_triggered = False
    st.session_state.file_uploaded = False
    st.session_state.word_count_uploaded = " "
    st.session_state.word_count_summary = " "
    st.session_state.downloadable_files = {}
    st.session_state.final_summary = " "        
    st.session_state.file = None
    st.session_state.text = " "     

# Run the main function
if __name__ == "__main__":
    main_function()
